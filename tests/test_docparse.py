"""docparse.py 단위 테스트 (문서 텍스트 추출)."""
import io

import pytest

import docparse


def test_txt():
    assert docparse.extract_text("a.txt", "현황 텍스트".encode("utf-8")) == "현황 텍스트"


def test_md():
    assert "제목" in docparse.extract_text("note.md", "# 제목".encode("utf-8"))


def test_unsupported_format():
    with pytest.raises(ValueError):
        docparse.extract_text("policy.hwp", b"x")


def test_docx_roundtrip():
    import docx

    document = docx.Document()
    document.add_paragraph("비밀번호 정책")
    document.add_paragraph("접근권한 반기 검토")
    buf = io.BytesIO()
    document.save(buf)

    out = docparse.extract_text("policy.docx", buf.getvalue())
    assert "비밀번호 정책" in out
    assert "접근권한 반기 검토" in out
